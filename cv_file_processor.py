import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import logging

try:
    import PyPDF2
except ImportError:
    print("PyPDF2 not installed. Install with: pip install PyPDF2")
    PyPDF2 = None

try:
    import docx
except ImportError:
    print("python-docx not installed. Install with: pip install python-docx")
    docx = None

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CVFileProcessor:
    """
    A class to handle CV file uploads and text extraction
    Supports PDF, DOCX, and TXT files
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def is_valid_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate file type and size
        
        Args:
            filename: Name of the uploaded file
            file_size: Size of the file in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not filename:
            return False, "No file provided"
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        # Check file size
        if file_size > self.MAX_FILE_SIZE:
            return False, f"File too large. Maximum size: {self.MAX_FILE_SIZE / (1024*1024):.1f}MB"
        
        return True, ""
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[bool, str, str]:
        """
        Extract text from PDF file using PyPDF2
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        if not PyPDF2:
            return False, "", "PyPDF2 library not installed"
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    return False, "", "PDF file appears to be empty"
                
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                            logger.info(f"Extracted text from page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text_content:
                    return False, "", "No readable text found in PDF"
                
                full_text = "\n\n".join(text_content)
                
                # Clean up the text
                full_text = self._clean_extracted_text(full_text)
                
                return True, full_text, ""
                
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return False, "", f"Error reading PDF: {str(e)}"
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[bool, str, str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        if not docx:
            return False, "", "python-docx library not installed"
        
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                return False, "", "No readable text found in DOCX file"
            
            full_text = "\n\n".join(text_content)
            full_text = self._clean_extracted_text(full_text)
            
            return True, full_text, ""
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return False, "", f"Error reading DOCX: {str(e)}"
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[bool, str, str]:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        if text.strip():
                            cleaned_text = self._clean_extracted_text(text)
                            return True, cleaned_text, ""
                except UnicodeDecodeError:
                    continue
            
            return False, "", "Could not decode text file with any supported encoding"
            
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return False, "", f"Error reading TXT file: {str(e)}"
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())  # Remove extra spaces
            if cleaned_line:  # Skip empty lines
                lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive line breaks (more than 2 consecutive)
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text.strip()
    
    def process_uploaded_file(self, file_data: bytes, filename: str) -> Tuple[bool, str, str]:
        """
        Process an uploaded file and extract text
        
        Args:
            file_data: Raw file data as bytes
            filename: Original filename
            
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        # Validate file
        is_valid, error_msg = self.is_valid_file(filename, len(file_data))
        if not is_valid:
            return False, "", error_msg
        
        # Create temporary file
        file_ext = Path(filename).suffix.lower()
        temp_file = None
        
        try:
            # Create temporary file with proper extension
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            logger.info(f"Processing {filename} ({len(file_data)} bytes)")
            
            # Extract text based on file type
            if file_ext == '.pdf':
                success, text, error = self.extract_text_from_pdf(temp_file_path)
            elif file_ext in ['.docx', '.doc']:
                success, text, error = self.extract_text_from_docx(temp_file_path)
            elif file_ext == '.txt':
                success, text, error = self.extract_text_from_txt(temp_file_path)
            else:
                return False, "", f"Unsupported file type: {file_ext}"
            
            if success:
                logger.info(f"Successfully extracted {len(text)} characters from {filename}")
                return True, text, ""
            else:
                return False, "", error
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return False, "", f"Error processing file: {str(e)}"
        
        finally:
            # Clean up temporary file
            if temp_file and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                    logger.debug(f"Cleaned up temporary file: {temp_file_path}")
                except Exception as e:
                    logger.warning(f"Could not clean up temporary file: {e}")

    def get_file_info(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Get information about the uploaded file
        
        Args:
            file_data: Raw file data as bytes
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        file_ext = Path(filename).suffix.lower() if filename else ""
        file_size = len(file_data)
        
        return {
            'filename': filename,
            'file_extension': file_ext,
            'file_size_bytes': file_size,
            'file_size_mb': round(file_size / (1024 * 1024), 2),
            'is_supported': file_ext in self.SUPPORTED_EXTENSIONS,
            'max_size_mb': self.MAX_FILE_SIZE / (1024 * 1024)
        }


# Example usage and testing
if __name__ == "__main__":
    processor = CVFileProcessor()
    
    # Test with a sample file (you would replace this with actual file upload handling)
    print("CV File Processor initialized")
    print(f"Supported extensions: {processor.SUPPORTED_EXTENSIONS}")
    print(f"Maximum file size: {processor.MAX_FILE_SIZE / (1024*1024):.1f}MB")
    
    # Example of how to use with Flask
    """
    # In your Flask route:
    
    from flask import Flask, request, jsonify
    
    app = Flask(__name__)
    cv_processor = CVFileProcessor()
    
    @app.route('/upload_cv', methods=['POST'])
    def upload_cv():
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        try:
            # Read file data
            file_data = file.read()
            filename = file.filename
            
            # Process the file
            success, extracted_text, error = cv_processor.process_uploaded_file(file_data, filename)
            
            if success:
                # Now you can process the extracted text with your existing RAG pipeline
                # For example: process_with_rag(extracted_text)
                
                return jsonify({
                    'success': True,
                    'extracted_text': extracted_text,
                    'file_info': cv_processor.get_file_info(file_data, filename)
                })
            else:
                return jsonify({'success': False, 'error': error})
                
        except Exception as e:
            return jsonify({'success': False, 'error': f'Server error: {str(e)}'})
    """